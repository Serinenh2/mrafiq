"""Fusion des modèles Word/PDF avec les informations de l'entreprise (§18/§41)."""
import io
import os
import shutil
from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Mm
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image as PILImage
from pypdf import PdfReader, PdfWriter
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfgen import canvas as pdfcanvas
from compliance.exports import ensure_arabic_fonts, ar_text

TOKEN = 'اسم الهيئة'

# Coordonnées (en points, repère PDF bas-gauche) des blancs du modèle
# « decision-designation-dpo.pdf » — relevées sur le gabarit source, article 1er.
# Chaque blanc est seul sur sa ligne : la largeur va jusqu'à la marge gauche (~70pt).
# (attribut Company, x du bord droit du blanc, y du bas du blanc, largeur max du blanc)
_DECISION_DPO_BOXES = [
    ('dpo_name', 391.4, 556.1, 321.4),
    ('dpo_poste', 390.9, 588.0, 320.9),
    ('dpo_phone', 409.0, 618.4, 339.0),
    ('dpo_email', 334.5, 648.6, 264.5),
]
_DECISION_COMPANY_BOX = (225.1, 679.0, 154.2)  # « ...لدى (nom de l'entreprise) »

def _fit_arabic_font_size(text, max_width, start=13, minimum=8):
    size = start
    while size > minimum and pdfmetrics.stringWidth(text, 'ArabicUI', size) > max_width:
        size -= 0.5
    return size

def _draw_decision_fields(c, page_h, company):
    """Superpose les informations personnelles du DPO, sa fonction et le nom de l'entreprise
    sur le modèle « Décision de désignation du DPO » (PDF non modifiable, blancs positionnés
    à partir des coordonnées relevées sur le gabarit source)."""
    ensure_arabic_fonts()
    for attr, x_right, y_bottom, max_w in _DECISION_DPO_BOXES:
        value = (getattr(company, attr, '') or '').strip()
        if not value:
            continue
        txt = ar_text(value)
        c.setFont('ArabicUI', _fit_arabic_font_size(txt, max_w))
        c.drawRightString(x_right, page_h - y_bottom + 5, txt)
    if company.name:
        x_right, y_bottom, max_w = _DECISION_COMPANY_BOX
        txt = ar_text(company.name)
        c.setFont('ArabicUI', _fit_arabic_font_size(txt, max_w))
        c.drawRightString(x_right, page_h - y_bottom + 5, txt)

def _replace_in_paragraph(paragraph, replacement):
    """Remplace le jeton même s'il est réparti entre plusieurs runs Word
    (fréquent avec le texte arabe) : on retravaille au niveau du paragraphe
    plutôt que run par run."""
    if TOKEN not in paragraph.text or not paragraph.runs:
        return
    new_text = paragraph.text.replace(TOKEN, replacement)
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ''

def _logo_path(company):
    logo = getattr(company, 'logo', None) if company else None
    if not logo:
        return None
    try:
        path = logo.path
        return path if os.path.exists(path) else None
    except Exception:
        return None

def _add_logo_to_docx(doc, company):
    """Insère le logo de l'entreprise en haut du document, s'il est défini (import logo)."""
    path = _logo_path(company)
    if not path:
        return
    try:
        anchor = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
        p = anchor.insert_paragraph_before()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(path, height=Mm(16))
    except Exception:
        pass

def merge_docx(source_path, company):
    """Retourne un buffer .docx avec le nom et le logo de l'entreprise injectés."""
    doc = DocxDocument(source_path)
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, company.name)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, company.name)
    _add_logo_to_docx(doc, company)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def stamp_docx(source_path, company):
    """Pour les modèles .docx non fusionnables (champs à compléter à la main) :
    texte inchangé, seul le logo de l'entreprise est ajouté."""
    doc = DocxDocument(source_path)
    _add_logo_to_docx(doc, company)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def stamp_pdf(source_path, company, extra_draw=None):
    """Superpose le logo de l'entreprise (et, si fourni, un dessin additionnel) sur la
    première page d'un PDF non modifiable."""
    logo_path = _logo_path(company)
    if not logo_path and not extra_draw:
        return copy_as_is(source_path)
    try:
        reader = PdfReader(source_path)
        page_w = float(reader.pages[0].mediabox.width)
        page_h = float(reader.pages[0].mediabox.height)

        overlay_buf = io.BytesIO()
        c = pdfcanvas.Canvas(overlay_buf, pagesize=(page_w, page_h))
        if logo_path:
            iw, ih = PILImage.open(logo_path).size
            max_w, max_h = 26 * mm, 12 * mm
            ratio = min(max_w / iw, max_h / ih)
            w, h = iw * ratio, ih * ratio
            c.drawImage(logo_path, page_w - 14 * mm - w, page_h - 14 * mm - h,
                        width=w, height=h, mask='auto', preserveAspectRatio=True)
        if extra_draw:
            extra_draw(c, page_h, company)
        c.save()
        overlay_buf.seek(0)
        overlay_page = PdfReader(overlay_buf).pages[0]

        writer = PdfWriter()
        for i, page in enumerate(reader.pages):
            if i == 0:
                page.merge_page(overlay_page)
            writer.add_page(page)
        out = io.BytesIO()
        writer.write(out)
        out.seek(0)
        return out
    except Exception:
        return copy_as_is(source_path)

def stamp_decision_pdf(source_path, company):
    """Modèle « Décision de désignation du DPO » : logo + informations personnelles du
    DPO (nom, téléphone, email) et nom de l'entreprise, superposés sur le PDF source."""
    return stamp_pdf(source_path, company, extra_draw=_draw_decision_fields)

def preview_payload(path):
    """Structure de blocs en lecture seule pour afficher un aperçu dans l'app (§18)."""
    path = str(path)
    if path.lower().endswith('.pdf'):
        return {'type': 'pdf'}
    doc = DocxDocument(path)
    blocks = []
    for child in doc.element.body.iterchildren():
        if child.tag == qn('w:p'):
            p = Paragraph(child, doc)
            if not p.text.strip():
                continue
            heading = p.style.name.startswith(('Heading', 'Title')) or \
                (bool(p.runs) and all(r.bold for r in p.runs))
            blocks.append({'type': 'heading' if heading else 'paragraph', 'text': p.text})
        elif child.tag == qn('w:tbl'):
            tbl = Table(child, doc)
            rows = [[cell.text for cell in row.cells] for row in tbl.rows]
            if any(any(c.strip() for c in r) for r in rows):
                blocks.append({'type': 'table', 'rows': rows})
    return {'type': 'docx', 'blocks': blocks}

def copy_as_is(source_path):
    """Copie brute (dernier recours si le logo ne peut pas être superposé)."""
    buf = io.BytesIO()
    with open(source_path, 'rb') as f:
        shutil.copyfileobj(f, buf)
    buf.seek(0)
    return buf
